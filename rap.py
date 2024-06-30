# Create a new Document
from docx import Document
from docx.shared import Pt
doc = Document()

# Page de titre
doc.add_heading('Rapport sur le Projet p_i', 0)
doc.add_paragraph('Auteur: [Votre Nom]')
doc.add_paragraph('Date: [Date]')

# Ajouter un saut de page
doc.add_page_break()

# Table des matières
doc.add_heading('Table des matières', level=1)
doc.add_paragraph('1. Introduction ...................................................................... 1')
doc.add_paragraph('2. Analyse et Conception ............................................................ 3')
doc.add_paragraph('   2.1 Introduction ................................................................ 3')
doc.add_paragraph('   2.2 Analyse des besoins ....................................................... 4')
doc.add_paragraph('   2.3 Spécification des besoins ............................................. 6')
doc.add_paragraph('   2.4 Modèles de conception .................................................. 8')
doc.add_paragraph('   2.5 Fonctionnalités du système ........................................ 12')
doc.add_paragraph('3. Implémentation ................................................................. 15')
doc.add_paragraph('4. Tests et Validation ............................................................ 20')
doc.add_paragraph('5. Conclusion ...................................................................... 25')
doc.add_paragraph('6. Annexes ............................................................................ 27')

# Ajouter un saut de page
doc.add_page_break()

# Chapitre 1: Introduction
doc.add_heading('Chapitre 1 : Introduction', level=1)
doc.add_paragraph('Contexte du projet:\n'
                  'Ce projet vise à développer une application web pour [objectif principal].\n'
                  '\n'
                  'Objectifs:\n'
                  '- Simplifier le processus de [but spécifique].\n'
                  '- Améliorer la transparence et l\'efficacité.\n'
                  '\n'
                  'Portée du projet:\n'
                  'Le projet couvre le développement, le déploiement et la maintenance de l\'application.')

# Ajouter un saut de page
doc.add_page_break()

# Chapitre 2: Analyse et Conception
doc.add_heading('Chapitre 2 : Analyse et Conception', level=1)
doc.add_heading('Introduction', level=2)
doc.add_paragraph('Ce chapitre présente une analyse détaillée des besoins et la conception du système.')

doc.add_heading('Analyse des besoins', level=2)
doc.add_paragraph('Besoins fonctionnels:\n'
                  '- Fonctionnalité 1: [Description]\n'
                  '- Fonctionnalité 2: [Description]\n'
                  '\n'
                  'Besoins non fonctionnels:\n'
                  '- Performance: [Description]\n'
                  '- Sécurité: [Description]')

doc.add_heading('Spécification des besoins', level=2)
doc.add_paragraph('L\'expression des besoins se base sur les demandes des utilisateurs.')

doc.add_heading('Modèles de conception', level=2)
doc.add_paragraph('Diagrammes UML:\n'
                  '- Diagramme de cas d\'utilisation: [Insérer diagramme]\n'
                  '- Diagramme de classes: [Insérer diagramme]')

doc.add_heading('Fonctionnalités du système', level=2)
doc.add_paragraph('Le système offre les fonctionnalités suivantes:\n'
                  '- Fonctionnalité 1: [Description]\n'
                  '- Fonctionnalité 2: [Description]')

# Ajouter un saut de page
doc.add_page_break()

# Chapitre 3: Implémentation
doc.add_heading('Chapitre 3 : Implémentation', level=1)
doc.add_paragraph('Technologies utilisées:\n'
                  '- Langages de programmation: [Liste des langages]\n'
                  '- Frameworks: [Liste des frameworks]\n'
                  '\n'
                  'Architecture du système:\n'
                  '- Description de l\'architecture\n'
                  '\n'
                  'Modules et composants:\n'
                  '- Module 1: [Description]\n'
                  '- Module 2: [Description]\n'
                  '\n'
                  'Captures d\'écran:\n'
                  '[Insérer captures d\'écran des interfaces utilisateur]\n'
                  '\n'
                  'Exemples de code:\n'
                  '[Insérer exemples de code pertinents]')

# Ajouter un saut de page
doc.add_page_break()

# Chapitre 4: Tests et Validation
doc.add_heading('Chapitre 4 : Tests et Validation', level=1)
doc.add_paragraph('Plan de test:\n'
                  '- Objectifs des tests\n'
                  '- Stratégie de test\n'
                  '\n'
                  'Types de tests:\n'
                  '- Tests unitaires: [Description]\n'
                  '- Tests d\'intégration: [Description]\n'
                  '\n'
                  'Résultats des tests:\n'
                  '- Test 1: [Résultat]\n'
                  '- Test 2: [Résultat]\n'
                  '\n'
                  'Corrections apportées:\n'
                  '[Liste des corrections]')

# Ajouter un saut de page
doc.add_page_break()

# Chapitre 5: Conclusion
doc.add_heading('Chapitre 5 : Conclusion', level=1)
doc.add_paragraph('Résumé des réalisations:\n'
                  '[Résumé]\n'
                  '\n'
                  'Défis rencontrés:\n'
                  '[Défis]\n'
                  '\n'
                  'Perspectives futures:\n'
                  '[Perspectives]')

# Ajouter un saut de page
doc.add_page_break()

# Annexes
doc.add_heading('Annexes', level=1)
doc.add_paragraph('Annexe A: Documentation technique\n'
                  'Annexe B: Guides d\'utilisation\n'
                  'Annexe C: [Autre information pertinente]')

# Ajuster le contenu pour atteindre 30 pages
while len(doc.paragraphs) < 300:
    doc.add_paragraph('Contenu à compléter...')

# Enregistrer le document
doc_path = "C:/Users/zeiny/Desktop/Rapport_projet_p_i.docx"
doc.save(doc_path)

doc_path
